from pathlib import Path


def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()

    if ext in {'.txt', '.md', '.csv', '.srt', '.log'}:
        # try utf-8, fallback to cp949 for Korean Windows files
        try:
            return file_path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            return file_path.read_text(encoding='cp949', errors='replace')

    if ext == '.docx':
        from docx import Document
        doc = Document(str(file_path))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        # tables
        for table in doc.tables:
            for row in table.rows:
                parts.append('\t'.join(cell.text for cell in row.cells))
        return '\n'.join(parts)

    # Unsupported: return a helpful message
    return f"[Unsupported file type: {ext}. Please upload a .txt or .docx file.]"
